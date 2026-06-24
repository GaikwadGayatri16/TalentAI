import os
import pdfplumber
import PyPDF2
from docx import Document

def parse_pdf(file_path: str) -> str:
    """
    Parses a PDF file using pdfplumber as the primary parser, 
    falling back to PyPDF2 if needed.
    """
    text = ""
    # Try pdfplumber first
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        print(f"pdfplumber failed for {file_path}, trying PyPDF2: {e}")
        text = ""

    # Fallback to PyPDF2 if pdfplumber failed or returned empty text
    if not text.strip():
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"PyPDF2 also failed for {file_path}: {e}")
            
    return text.strip()

def parse_docx(file_path: str) -> str:
    """
    Parses a DOCX file using python-docx.
    """
    text = ""
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells if cell.text]
                if row_text:
                    text += " | ".join(row_text) + "\n"
    except Exception as e:
        print(f"python-docx failed for {file_path}: {e}")
    return text.strip()

def parse_resume(file_path: str) -> str:
    """
    Determines file extension and calls the appropriate parser.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
        
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext in [".docx", ".doc"]:
        return parse_docx(file_path)
    else:
        # Fallback to reading file directly as text
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            print(f"Direct text read failed for {file_path}: {e}")
            return ""
